import random
import socket
import PacketTCP
from docx import Document
import threading

DATA_MAX_LEN = 194


def read_docx():
    filepath = "test.docx"
    doc = Document(filepath)
    # 初始化一个空字符串
    full_text = []
    # 遍历所有段落
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 将段落列表连接成一个字符串
    test1 = '\n'.join(full_text)
    test2 = ''
    for _ in range(0, len(test1)):
        if test1[_] == ' ':
            test2 += '$'
        else:
            test2 += test1[_]
    return test2


def write_docx(text2):
    test1 = ''
    for _ in range(0, len(text2)):
        if text2[_] == '$':
            test1 += ' '
        else:
            test1 += text2[_]
    filepath = "test_new.docx"
    doc = Document()
    for line in test1.split('\n'):
        doc.add_paragraph(line)
    doc.save(filepath)
    print('保存完毕')


def Init_And_Agree():
    lst = []
    global pointer, border
    lev_len = border - pointer
    while 1:
        Size = int(random.random() * lev_len + 0.1)
        if Size < Lmin:
            continue

        if lev_len - Size < 2:
            Size = lev_len
        elif Size == 0 and lev_len != 0:
            continue
        break

    num, now, Rborder = 0, pointer, pointer + Size - 1
    choiceLst = [_ for _ in range(Lmin, Lmax + 1)]
    while now <= Rborder:
        choiceSize = random.choice(choiceLst)
        if now + choiceSize - 1 > Rborder:
            num += 1
            lst.append(test[now:Rborder + 1])
            break
        else:
            num += 1
            lst.append(test[now:now + choiceSize])
            now = now + choiceSize

    while 1:
        # 发送Initialization报文
        data = PacketTCP.Initialization(1, num)
        clientsocket.send(data.pack())

        # 接收Agree报文
        newdata = clientsocket.recv(1024)
        if not newdata:
            break
        packet_received = PacketTCP.Agree.unpack(newdata)
        print('Received from server: ', packet_received.callme())
        if packet_received.Type == 2:
            break

    return Size, lst


def request_and_reverse(test_to_send):
    while 1:
        # 发送ReverseRequest报文
        data = PacketTCP.ReverseRequest(3, len(test_to_send), test_to_send)
        clientsocket.send(data.pack())
        # 接收ReverseAnswer报文
        newdata = clientsocket.recv(1024)
        packet_received = PacketTCP.ReverseAnswer.unpack(newdata)
        packet_received.reverseData = packet_received.reverseData.strip('\n ')
        print('Received from server: ', packet_received.callme())
        if packet_received.Type == 4:
            # print("发送与接收: ", '\n', data.Data, '\n', packet_received.reverseData)
            break
    return packet_received.reverseData


def check_flag():
    global flagcontinue
    while flagcontinue:
        s = input()
        if s == "end":
            flagcontinue = False


def Input_Lmin_Lmax():
    global Lmin, Lmax, pointer, border
    while True:
        try:
            print(f"当前剩余文本长度: {border - pointer}  最大数据的长度为:{DATA_MAX_LEN}")
            Lmin = int(input("请输入Lmin: "))
            if Lmin > border - pointer:
                print("Lmin 应该小于等于剩余文本长度")
                continue

            if Lmin > DATA_MAX_LEN:
                print("Lmin 应该小于等于最大数据的长度")
                continue

            Lmax = int(input("请输入Lmax: "))
            if Lmin > Lmax:
                print("Lmin 必须小于等于 Lmax")
                continue

            if Lmax > DATA_MAX_LEN:
                print("Lmax 应该小于等于最大数据的长度")
                continue

            break
        except Exception as e1:
            print("请输入整数!")


if __name__ == "__main__":
    # flagcontinue = True
    # spe = threading.Thread(target=check_flag)
    # spe.start()
    ipaddress = input('请输入ip地址：')
    portaddress = input('请输入端口号：')
    test = read_docx()
    clientsocket = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
    clientsocket.connect((ipaddress, int(portaddress)))
    # clientsocket.connect(('localhost', 1118))
    border = len(test)
    pointer = 0
    ReverseTest = ""
    Lmin, Lmax = 16, 16
    while pointer < border:
        # 输入Lmin,Lmax
        Input_Lmin_Lmax()

        blocks_size, to_reverse_tests = Init_And_Agree()
        # print(blocks_size, *to_reverse_tests)
        for i in range(0, len(to_reverse_tests)):
            temp = request_and_reverse(to_reverse_tests[i]).strip('\n ')
            out_temp = ""
            for _ in temp:
                if _ == '$':
                    out_temp += ' '
                else:
                    out_temp += _
            print(f"{i + 1}: {out_temp}")
            ReverseTest += temp
        pointer += blocks_size
    write_docx(ReverseTest)
    # print(ReverseTest)
    # while flagcontinue:
    #     pass
    clientsocket.close()
